"""
Document Generation Service
Generates branded Leadway Health renewal notices (.docx + PDF).
ONLY called for APPROVED policies (enforcement is in the API layer).
"""
import os, subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.models.renewal import RenewalPolicy, RenewalStatus
from app.core.config import settings
import logging

logger = logging.getLogger(__name__)

BRAND_BLUE = RGBColor(0x00, 0x2F, 0x6C)
BRAND_RED  = RGBColor(0xE3, 0x06, 0x13)
BRAND_GRAY = RGBColor(0x60, 0x60, 0x60)


def _table_row(table, label: str, value: str, highlight: bool = False):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    runs0 = row.cells[0].paragraphs[0].runs
    runs1 = row.cells[1].paragraphs[0].runs
    if runs0:
        runs0[0].bold = True
        runs0[0].font.color.rgb = BRAND_GRAY
    if runs1 and highlight:
        runs1[0].font.color.rgb = BRAND_RED
        runs1[0].bold = True


def generate_renewal_notice(policy: RenewalPolicy, db: Session) -> dict:
    """
    Produce a .docx renewal notice for an APPROVED policy.
    Includes SAME RATE wording for COR < 80%.
    Lists the approving authority where applicable.
    """
    os.makedirs(settings.DOCUMENTS_DIR, exist_ok=True)
    safe = "".join(c for c in policy.company_name if c.isalnum() or c in " _-").strip()
    base = f"renewal_{policy.policy_number}_{safe}".replace(" ", "_")
    docx_path = os.path.join(settings.DOCUMENTS_DIR, f"{base}.docx")
    pdf_path  = os.path.join(settings.DOCUMENTS_DIR, f"{base}.pdf")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin = sec.right_margin = Cm(2.5)

    # Header
    hdr = sec.header.paragraphs[0]
    hdr.clear()
    r = hdr.add_run("LEADWAY HEALTH INSURANCE LIMITED")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BRAND_BLUE
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = hdr.add_run("\nMonthly Renewal Automation System (MRAS)")
    sub.font.size = Pt(9); sub.font.color.rgb = BRAND_GRAY

    # Title
    t = doc.add_paragraph()
    tr = t.add_run("POLICY RENEWAL NOTICE")
    tr.bold = True; tr.font.size = Pt(17); tr.font.color.rgb = BRAND_BLUE
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle: same-rate wording
    cor_pct = (policy.cor or 0) * 100
    if cor_pct < 80:
        sub_text = "RATE MAINTAINED — No adjustment required based on current loss experience"
        sub_color = RGBColor(0x15, 0x80, 0x3D)
    else:
        sub_text = f"RATE ADJUSTMENT NOTICE — {policy.renewal_rate:.1f}% increase applied"
        sub_color = BRAND_RED
    s = doc.add_paragraph()
    sr = s.add_run(sub_text)
    sr.font.size = Pt(10); sr.font.color.rgb = sub_color; sr.bold = True
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.runs[0].font.size = Pt(9); d.runs[0].font.color.rgb = BRAND_GRAY
    doc.add_paragraph()

    # Sections
    def section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND_BLUE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def new_table():
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(2.6)
        tbl.columns[1].width = Inches(4.0)
        return tbl

    section_heading("CLIENT INFORMATION")
    ct = new_table()
    _table_row(ct, "Company",        policy.company_name)
    _table_row(ct, "Policy Number",  policy.policy_number)
    if policy.scheme_ref:
        _table_row(ct, "Scheme Ref", policy.scheme_ref)
    _table_row(ct, "Segment",        policy.segment.value)
    if policy.business_sector:
        _table_row(ct, "Sector",     policy.business_sector)
    _table_row(ct, "Lives Covered",  str(policy.no_of_lives or "—"))
    _table_row(ct, "Contact Person", policy.contact_name or "—")
    _table_row(ct, "Email",          policy.contact_email or "—")

    doc.add_paragraph()
    section_heading("POLICY DATES")
    dt = new_table()
    _table_row(dt, "Start Date",     policy.start_date.strftime("%d %B %Y") if policy.start_date else "—")
    _table_row(dt, "Renewal Date",   policy.end_date.strftime("%d %B %Y"))
    if policy.is_pro_rata:
        _table_row(dt, "Policy Period", f"{policy.policy_months:.1f} months (pro-rata)", highlight=True)
    _table_row(dt, "Days to Renewal", str(policy.days_to_renewal) if policy.days_to_renewal is not None else "—")

    doc.add_paragraph()
    section_heading("FINANCIAL METRICS")
    ft = new_table()
    _table_row(ft, "Current Annual Premium",  f"₦{policy.current_premium:,.2f}")
    _table_row(ft, "Written/Earned Premium",  f"₦{policy.earned_premium or policy.total_premium:,.2f}")
    _table_row(ft, "Total Claims",            f"₦{policy.total_claims:,.2f}")
    _table_row(ft, "Loss Ratio (LR)",         f"{cor_pct - 15:.1f}%",  highlight=cor_pct >= 95)
    _table_row(ft, "Combined Ratio (COR)",    f"{cor_pct:.1f}%",       highlight=cor_pct >= 95)

    doc.add_paragraph()
    section_heading("RENEWAL TERMS")
    rt = new_table()
    _table_row(rt, "Rate Band",               policy.rate_band or "—")
    _table_row(rt, "Rate Adjustment",
               "0% — Same Rate" if cor_pct < 80 else f"{policy.renewal_rate:.1f}%")
    _table_row(rt, "Proposed Renewal Premium", f"₦{policy.renewal_premium:,.2f}", highlight=True)

    # Approving authority
    if policy.md_ceo_approved_at:
        authority = "MD/CEO"
    elif policy.hbd_approved_at:
        authority = "Head of Business Development"
    elif policy.underwriter_approved_at:
        authority = "Underwriter"
    elif policy.sales_confirmed_at:
        authority = "Sales Officer"
    else:
        authority = "Automated (No Approval Required)"
    _table_row(rt, "Approving Authority",     authority)
    if policy.sales_confirmed_at:
        _table_row(rt, "Sales Confirmed At",  policy.sales_confirmed_at.strftime("%d %B %Y %H:%M"))
    if policy.hbd_approved_at:
        _table_row(rt, "HBD Approved At",     policy.hbd_approved_at.strftime("%d %B %Y %H:%M"))
    if policy.md_ceo_approved_at:
        _table_row(rt, "MD/CEO Approved At",  policy.md_ceo_approved_at.strftime("%d %B %Y %H:%M"))

    doc.add_paragraph()
    note = doc.add_paragraph(
        "This document is system-generated by the Leadway MRAS. "
        "Figures are subject to final underwriting review. "
        "For queries, contact your relationship manager."
    )
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.color.rgb = BRAND_GRAY
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(docx_path)
    logger.info(f"DOCX saved: {docx_path}")

    pdf_converted = False
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", settings.DOCUMENTS_DIR, docx_path],
            capture_output=True, timeout=30,
        )
        pdf_converted = result.returncode == 0
    except Exception as e:
        logger.warning(f"LibreOffice PDF conversion failed: {e}")

    return {"docx_path": docx_path, "pdf_path": pdf_path if pdf_converted else None}
